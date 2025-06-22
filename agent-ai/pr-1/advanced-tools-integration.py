# LocalAgentSystem - Advanced Tools & Integration
# 高级工具集和集成功能

import asyncio
import aiohttp
from typing import Dict, List, Any, Optional, Union
import pandas as pd
import numpy as np
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker
import pymongo
from elasticsearch import AsyncElasticsearch
import docker
import kubernetes
from jinja2 import Template
import markdown
import pdfkit
from docx import Document
import xlsxwriter
import csv
import json
import yaml
import toml
from pathlib import Path
import hashlib
import jwt
from cryptography.fernet import Fernet
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders


# ==================== 数据库工具 ====================

class DatabaseTool(Tool):
    """数据库操作工具 - 支持多种数据库"""
    
    def __init__(self):
        super().__init__(
            name="database_tool",
            description="Execute queries and manage databases (PostgreSQL, MySQL, MongoDB, etc.)"
        )
        self.connections = {}
        
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行数据库操作"""
        try:
            if action == "connect":
                return await self.connect_database(**kwargs)
            elif action == "query":
                return await self.execute_query(**kwargs)
            elif action == "insert":
                return await self.insert_data(**kwargs)
            elif action == "update":
                return await self.update_data(**kwargs)
            elif action == "delete":
                return await self.delete_data(**kwargs)
            elif action == "create_table":
                return await self.create_table(**kwargs)
            elif action == "migrate":
                return await self.run_migration(**kwargs)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def connect_database(self, db_type: str, connection_string: str, 
                              alias: str = "default") -> Dict[str, Any]:
        """连接数据库"""
        try:
            if db_type in ["postgresql", "mysql", "sqlite"]:
                engine = create_engine(connection_string)
                Session = sessionmaker(bind=engine)
                self.connections[alias] = {
                    "type": db_type,
                    "engine": engine,
                    "session": Session
                }
            elif db_type == "mongodb":
                client = pymongo.MongoClient(connection_string)
                self.connections[alias] = {
                    "type": db_type,
                    "client": client
                }
            elif db_type == "elasticsearch":
                client = AsyncElasticsearch([connection_string])
                self.connections[alias] = {
                    "type": db_type,
                    "client": client
                }
            else:
                return {"success": False, "error": f"Unsupported database type: {db_type}"}
                
            return {"success": True, "message": f"Connected to {db_type} as '{alias}'"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def execute_query(self, query: str, alias: str = "default", 
                          params: Dict[str, Any] = None) -> Dict[str, Any]:
        """执行查询"""
        if alias not in self.connections:
            return {"success": False, "error": f"No connection found for alias: {alias}"}
            
        conn_info = self.connections[alias]
        db_type = conn_info["type"]
        
        try:
            if db_type in ["postgresql", "mysql", "sqlite"]:
                with conn_info["engine"].connect() as conn:
                    result = conn.execute(text(query), params or {})
                    if result.returns_rows:
                        rows = result.fetchall()
                        columns = result.keys()
                        return {
                            "success": True,
                            "data": [dict(zip(columns, row)) for row in rows],
                            "row_count": len(rows)
                        }
                    else:
                        return {
                            "success": True,
                            "affected_rows": result.rowcount
                        }
                        
            elif db_type == "mongodb":
                # MongoDB查询处理
                db_name, collection_name = query.split(".")[:2]
                db = conn_info["client"][db_name]
                collection = db[collection_name]
                
                # 简单查询示例
                results = list(collection.find(params or {}))
                return {
                    "success": True,
                    "data": results,
                    "row_count": len(results)
                }
                
            elif db_type == "elasticsearch":
                # Elasticsearch查询处理
                response = await conn_info["client"].search(
                    index=params.get("index", "*"),
                    body=json.loads(query) if isinstance(query, str) else query
                )
                return {
                    "success": True,
                    "data": response["hits"]["hits"],
                    "total": response["hits"]["total"]["value"]
                }
                
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def create_table(self, table_name: str, schema: Dict[str, Any], 
                         alias: str = "default") -> Dict[str, Any]:
        """创建表"""
        if alias not in self.connections:
            return {"success": False, "error": f"No connection found for alias: {alias}"}
            
        conn_info = self.connections[alias]
        db_type = conn_info["type"]
        
        try:
            if db_type in ["postgresql", "mysql", "sqlite"]:
                # 构建CREATE TABLE语句
                columns = []
                for col_name, col_def in schema.items():
                    col_type = col_def.get("type", "VARCHAR(255)")
                    nullable = "" if col_def.get("nullable", True) else "NOT NULL"
                    primary = "PRIMARY KEY" if col_def.get("primary_key", False) else ""
                    columns.append(f"{col_name} {col_type} {nullable} {primary}")
                    
                create_sql = f"CREATE TABLE IF NOT EXISTS {table_name} ({', '.join(columns)})"
                
                with conn_info["engine"].connect() as conn:
                    conn.execute(text(create_sql))
                    conn.commit()
                    
                return {"success": True, "message": f"Table {table_name} created"}
                
            else:
                return {"success": False, "error": f"Create table not supported for {db_type}"}
                
        except Exception as e:
            return {"success": False, "error": str(e)}


# ==================== API集成工具 ====================

class APIIntegrationTool(Tool):
    """API集成工具 - 调用外部API"""
    
    def __init__(self):
        super().__init__(
            name="api_integration",
            description="Integrate with external APIs (REST, GraphQL, WebSocket)"
        )
        self.session = None
        self.api_configs = {}
        
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行API操作"""
        if not self.session:
            self.session = aiohttp.ClientSession()
            
        try:
            if action == "configure":
                return self.configure_api(**kwargs)
            elif action == "rest_call":
                return await self.make_rest_call(**kwargs)
            elif action == "graphql_query":
                return await self.make_graphql_query(**kwargs)
            elif action == "webhook":
                return await self.setup_webhook(**kwargs)
            elif action == "batch_request":
                return await self.batch_api_requests(**kwargs)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def configure_api(self, name: str, base_url: str, headers: Dict[str, str] = None,
                     auth_type: str = None, auth_config: Dict[str, Any] = None) -> Dict[str, Any]:
        """配置API"""
        self.api_configs[name] = {
            "base_url": base_url,
            "headers": headers or {},
            "auth_type": auth_type,
            "auth_config": auth_config or {}
        }
        
        return {"success": True, "message": f"API '{name}' configured"}
        
    async def make_rest_call(self, api_name: str = None, method: str = "GET",
                           endpoint: str = "", params: Dict[str, Any] = None,
                           data: Any = None, headers: Dict[str, str] = None) -> Dict[str, Any]:
        """执行REST API调用"""
        # 获取API配置
        if api_name and api_name in self.api_configs:
            config = self.api_configs[api_name]
            url = config["base_url"] + endpoint
            headers = {**config["headers"], **(headers or {})}
            
            # 处理认证
            if config["auth_type"] == "bearer":
                headers["Authorization"] = f"Bearer {config['auth_config']['token']}"
            elif config["auth_type"] == "api_key":
                headers[config['auth_config']['header']] = config['auth_config']['key']
        else:
            url = endpoint
            headers = headers or {}
            
        # 执行请求
        try:
            async with self.session.request(
                method=method,
                url=url,
                params=params,
                json=data if isinstance(data, (dict, list)) else None,
                data=data if isinstance(data, str) else None,
                headers=headers
            ) as response:
                response_data = await response.text()
                
                # 尝试解析JSON
                try:
                    response_data = json.loads(response_data)
                except:
                    pass
                    
                return {
                    "success": response.status < 400,
                    "status_code": response.status,
                    "headers": dict(response.headers),
                    "data": response_data
                }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def make_graphql_query(self, endpoint: str, query: str, 
                                variables: Dict[str, Any] = None) -> Dict[str, Any]:
        """执行GraphQL查询"""
        try:
            async with self.session.post(
                endpoint,
                json={
                    "query": query,
                    "variables": variables or {}
                },
                headers={"Content-Type": "application/json"}
            ) as response:
                data = await response.json()
                
                return {
                    "success": "errors" not in data,
                    "data": data.get("data"),
                    "errors": data.get("errors", [])
                }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def batch_api_requests(self, requests: List[Dict[str, Any]], 
                                max_concurrent: int = 5) -> Dict[str, Any]:
        """批量API请求"""
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def make_request_with_semaphore(request_config):
            async with semaphore:
                return await self.make_rest_call(**request_config)
                
        results = await asyncio.gather(*[
            make_request_with_semaphore(req) for req in requests
        ], return_exceptions=True)
        
        return {
            "success": True,
            "results": results,
            "total": len(results),
            "successful": sum(1 for r in results if isinstance(r, dict) and r.get("success"))
        }
        
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """清理资源"""
        if self.session:
            await self.session.close()


# ==================== 文档生成工具 ====================

class DocumentGeneratorTool(Tool):
    """文档生成工具 - 生成各种格式的文档"""
    
    def __init__(self):
        super().__init__(
            name="document_generator",
            description="Generate documents in various formats (PDF, Word, Excel, Markdown, etc.)"
        )
        self.templates = {}
        
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行文档生成操作"""
        try:
            if action == "generate_pdf":
                return await self.generate_pdf(**kwargs)
            elif action == "generate_word":
                return await self.generate_word(**kwargs)
            elif action == "generate_excel":
                return await self.generate_excel(**kwargs)
            elif action == "generate_markdown":
                return await self.generate_markdown(**kwargs)
            elif action == "generate_html":
                return await self.generate_html(**kwargs)
            elif action == "template_render":
                return await self.render_template(**kwargs)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def generate_pdf(self, content: str, output_path: str, 
                          format_type: str = "markdown") -> Dict[str, Any]:
        """生成PDF文档"""
        try:
            if format_type == "markdown":
                # Markdown转HTML
                html_content = markdown.markdown(content, extensions=['extra', 'codehilite'])
                html_content = f"""
                <html>
                <head>
                    <meta charset="utf-8">
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                        code {{ background-color: #f4f4f4; padding: 2px 4px; }}
                        pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; }}
                    </style>
                </head>
                <body>{html_content}</body>
                </html>
                """
            else:
                html_content = content
                
            # 生成PDF
            pdfkit.from_string(html_content, output_path)
            
            return {
                "success": True,
                "output_path": output_path,
                "message": "PDF generated successfully"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def generate_word(self, content: Union[str, List[Dict[str, Any]]], 
                          output_path: str) -> Dict[str, Any]:
        """生成Word文档"""
        try:
            doc = Document()
            
            if isinstance(content, str):
                # 简单文本
                doc.add_paragraph(content)
            else:
                # 结构化内容
                for section in content:
                    if section.get("type") == "heading":
                        doc.add_heading(section["text"], level=section.get("level", 1))
                    elif section.get("type") == "paragraph":
                        doc.add_paragraph(section["text"])
                    elif section.get("type") == "list":
                        for item in section["items"]:
                            doc.add_paragraph(item, style='List Bullet')
                    elif section.get("type") == "table":
                        table = doc.add_table(rows=len(section["data"]) + 1, 
                                            cols=len(section["headers"]))
                        # 添加表头
                        for i, header in enumerate(section["headers"]):
                            table.cell(0, i).text = header
                        # 添加数据
                        for row_idx, row_data in enumerate(section["data"]):
                            for col_idx, cell_data in enumerate(row_data):
                                table.cell(row_idx + 1, col_idx).text = str(cell_data)
                                
            doc.save(output_path)
            
            return {
                "success": True,
                "output_path": output_path,
                "message": "Word document generated successfully"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def generate_excel(self, data: Union[pd.DataFrame, List[Dict[str, Any]]], 
                           output_path: str, sheet_name: str = "Sheet1") -> Dict[str, Any]:
        """生成Excel文档"""
        try:
            # 转换为DataFrame
            if isinstance(data, list):
                df = pd.DataFrame(data)
            else:
                df = data
                
            # 创建Excel writer
            with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                # 获取工作簿和工作表
                workbook = writer.book
                worksheet = writer.sheets[sheet_name]
                
                # 添加格式
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BD',
                    'border': 1
                })
                
                # 应用格式到表头
                for col_num, value in enumerate(df.columns.values):
                    worksheet.write(0, col_num, value, header_format)
                    
                # 自动调整列宽
                for i, col in enumerate(df.columns):
                    column_len = df[col].astype(str).str.len().max()
                    column_len = max(column_len, len(col)) + 2
                    worksheet.set_column(i, i, column_len)
                    
            return {
                "success": True,
                "output_path": output_path,
                "message": "Excel file generated successfully",
                "rows": len(df),
                "columns": len(df.columns)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def render_template(self, template_name: str, context: Dict[str, Any],
                            output_path: str = None) -> Dict[str, Any]:
        """渲染模板"""
        try:
            if template_name not in self.templates:
                # 尝试从文件加载模板
                template_path = Path(f"templates/{template_name}")
                if template_path.exists():
                    with open(template_path, 'r') as f:
                        template_content = f.read()
                    self.templates[template_name] = Template(template_content)
                else:
                    return {"success": False, "error": f"Template '{template_name}' not found"}
                    
            template = self.templates[template_name]
            rendered = template.render(**context)
            
            if output_path:
                with open(output_path, 'w') as f:
                    f.write(rendered)
                    
            return {
                "success": True,
                "content": rendered,
                "output_path": output_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}


# ==================== 容器和编排工具 ====================

class ContainerOrchestrationTool(Tool):
    """容器编排工具 - Docker和Kubernetes管理"""
    
    def __init__(self):
        super().__init__(
            name="container_orchestration",
            description="Manage Docker containers and Kubernetes deployments"
        )
        self.docker_client = None
        self.k8s_client = None
        
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行容器操作"""
        try:
            if action.startswith("docker_"):
                return await self._execute_docker_action(action, **kwargs)
            elif action.startswith("k8s_"):
                return await self._execute_k8s_action(action, **kwargs)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def _execute_docker_action(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行Docker操作"""
        if not self.docker_client:
            self.docker_client = docker.from_env()
            
        try:
            if action == "docker_run":
                container = self.docker_client.containers.run(
                    kwargs.get("image"),
                    kwargs.get("command", ""),
                    detach=kwargs.get("detach", True),
                    ports=kwargs.get("ports", {}),
                    environment=kwargs.get("environment", {}),
                    volumes=kwargs.get("volumes", {}),
                    name=kwargs.get("name"),
                    remove=kwargs.get("remove", False)
                )
                return {
                    "success": True,
                    "container_id": container.id,
                    "status": container.status
                }
                
            elif action == "docker_list":
                containers = self.docker_client.containers.list(all=kwargs.get("all", False))
                return {
                    "success": True,
                    "containers": [
                        {
                            "id": c.id[:12],
                            "name": c.name,
                            "image": c.image.tags[0] if c.image.tags else "unknown",
                            "status": c.status,
                            "created": c.attrs["Created"]
                        }
                        for c in containers
                    ]
                }
                
            elif action == "docker_stop":
                container = self.docker_client.containers.get(kwargs.get("container_id"))
                container.stop()
                return {"success": True, "message": f"Container {container.id[:12]} stopped"}
                
            elif action == "docker_logs":
                container = self.docker_client.containers.get(kwargs.get("container_id"))
                logs = container.logs(tail=kwargs.get("lines", 100)).decode('utf-8')
                return {"success": True, "logs": logs}
                
            elif action == "docker_build":
                image, logs = self.docker_client.images.build(
                    path=kwargs.get("path", "."),
                    tag=kwargs.get("tag"),
                    dockerfile=kwargs.get("dockerfile", "Dockerfile")
                )
                return {
                    "success": True,
                    "image_id": image.id,
                    "tags": image.tags
                }
                
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def _execute_k8s_action(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行Kubernetes操作"""
        if not self.k8s_client:
            kubernetes.config.load_incluster_config()  # 或 load_kube_config()
            self.k8s_client = kubernetes.client
            
        try:
            if action == "k8s_deploy":
                # 部署应用
                apps_v1 = self.k8s_client.AppsV1Api()
                deployment = self._create_deployment_object(**kwargs)
                
                response = apps_v1.create_namespaced_deployment(
                    namespace=kwargs.get("namespace", "default"),
                    body=deployment
                )
                
                return {
                    "success": True,
                    "deployment_name": response.metadata.name,
                    "namespace": response.metadata.namespace
                }
                
            elif action == "k8s_list_pods":
                core_v1 = self.k8s_client.CoreV1Api()
                pods = core_v1.list_namespaced_pod(
                    namespace=kwargs.get("namespace", "default")
                )
                
                return {
                    "success": True,
                    "pods": [
                        {
                            "name": pod.metadata.name,
                            "status": pod.status.phase,
                            "ip": pod.status.pod_ip,
                            "node": pod.spec.node_name
                        }
                        for pod in pods.items
                    ]
                }
                
            elif action == "k8s_scale":
                apps_v1 = self.k8s_client.AppsV1Api()
                
                # 更新副本数
                body = {"spec": {"replicas": kwargs.get("replicas", 1)}}
                
                response = apps_v1.patch_namespaced_deployment_scale(
                    name=kwargs.get("deployment_name"),
                    namespace=kwargs.get("namespace", "default"),
                    body=body
                )
                
                return {
                    "success": True,
                    "deployment": response.metadata.name,
                    "replicas": response.spec.replicas
                }
                
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def _create_deployment_object(self, **kwargs):
        """创建Kubernetes部署对象"""
        # 简化的部署对象创建
        deployment = kubernetes.client.V1Deployment(
            api_version="apps/v1",
            kind="Deployment",
            metadata=kubernetes.client.V1ObjectMeta(name=kwargs.get("name")),
            spec=kubernetes.client.V1DeploymentSpec(
                replicas=kwargs.get("replicas", 1),
                selector=kubernetes.client.V1LabelSelector(
                    match_labels={"app": kwargs.get("name")}
                ),
                template=kubernetes.client.V1PodTemplateSpec(
                    metadata=kubernetes.client.V1ObjectMeta(
                        labels={"app": kwargs.get("name")}
                    ),
                    spec=kubernetes.client.V1PodSpec(
                        containers=[
                            kubernetes.client.V1Container(
                                name=kwargs.get("name"),
                                image=kwargs.get("image"),
                                ports=[
                                    kubernetes.client.V1ContainerPort(container_port=port)
                                    for port in kwargs.get("ports", [])
                                ]
                            )
                        ]
                    )
                )
            )
        )
        return deployment


# ==================== 数据处理工具 ====================

class DataProcessingTool(Tool):
    """数据处理工具 - 数据清洗、转换、分析"""
    
    def __init__(self):
        super().__init__(
            name="data_processing",
            description="Process, clean, transform, and analyze data"
        )
        
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行数据处理操作"""
        try:
            if action == "clean_data":
                return await self.clean_data(**kwargs)
            elif action == "transform_data":
                return await self.transform_data(**kwargs)
            elif action == "analyze_data":
                return await self.analyze_data(**kwargs)
            elif action == "merge_datasets":
                return await self.merge_datasets(**kwargs)
            elif action == "aggregate_data":
                return await self.aggregate_data(**kwargs)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def clean_data(self, data: Union[pd.DataFrame, List[Dict[str, Any]]], 
                        operations: List[str]) -> Dict[str, Any]:
        """清洗数据"""
        try:
            # 转换为DataFrame
            if isinstance(data, list):
                df = pd.DataFrame(data)
            else:
                df = data.copy()
                
            original_shape = df.shape
            
            for operation in operations:
                if operation == "remove_duplicates":
                    df = df.drop_duplicates()
                elif operation == "remove_nulls":
                    df = df.dropna()
                elif operation == "fill_nulls":
                    df = df.fillna(method='ffill').fillna(method='bfill')
                elif operation == "standardize_columns":
                    df.columns = df.columns.str.lower().str.replace(' ', '_')
                elif operation == "remove_outliers":
                    # 使用IQR方法移除异常值
                    Q1 = df.select_dtypes(include=[np.number]).quantile(0.25)
                    Q3 = df.select_dtypes(include=[np.number]).quantile(0.75)
                    IQR = Q3 - Q1
                    df = df[~((df < (Q1 - 1.5 * IQR)) | (df > (Q3 + 1.5 * IQR))).any(axis=1)]
                    
            return {
                "success": True,
                "cleaned_data": df.to_dict('records'),
                "original_shape": original_shape,
                "cleaned_shape": df.shape,
                "rows_removed": original_shape[0] - df.shape[0]
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def transform_data(self, data: pd.DataFrame, transformations: Dict[str, Any]) -> Dict[str, Any]:
        """转换数据"""
        try:
            df = data.copy()
            
            for column, transform in transformations.items():
                if transform["type"] == "normalize":
                    df[column] = (df[column] - df[column].min()) / (df[column].max() - df[column].min())
                elif transform["type"] == "standardize":
                    df[column] = (df[column] - df[column].mean()) / df[column].std()
                elif transform["type"] == "log":
                    df[column] = np.log1p(df[column])
                elif transform["type"] == "categorical":
                    df[column] = pd.Categorical(df[column]).codes
                elif transform["type"] == "datetime":
                    df[column] = pd.to_datetime(df[column])
                elif transform["type"] == "custom":
                    # 应用自定义函数
                    df[column] = df[column].apply(eval(transform["function"]))
                    
            return {
                "success": True,
                "transformed_data": df.to_dict('records'),
                "shape": df.shape,
                "dtypes": df.dtypes.to_dict()
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def analyze_data(self, data: pd.DataFrame) -> Dict[str, Any]:
        """分析数据"""
        try:
            analysis = {
                "shape": data.shape,
                "columns": list(data.columns),
                "dtypes": data.dtypes.to_dict(),
                "null_counts": data.isnull().sum().to_dict(),
                "summary_stats": data.describe().to_dict(),
                "unique_counts": {col: data[col].nunique() for col in data.columns},
                "correlations": data.select_dtypes(include=[np.number]).corr().to_dict()
            }
            
            # 检测数据质量问题
            issues = []
            
            # 检查空值
            null_cols = [col for col, count in analysis["null_counts"].items() if count > 0]
            if null_cols:
                issues.append(f"Columns with null values: {', '.join(null_cols)}")
                
            # 检查重复行
            duplicates = data.duplicated().sum()
            if duplicates > 0:
                issues.append(f"Found {duplicates} duplicate rows")
                
            # 检查高基数分类变量
            for col in data.select_dtypes(include=['object']).columns:
                if data[col].nunique() > data.shape[0] * 0.5:
                    issues.append(f"Column '{col}' has high cardinality")
                    
            analysis["data_quality_issues"] = issues
            
            return {
                "success": True,
                "analysis": analysis
            }
        except Exception as e:
            return {"success": False, "error": str(e)}


# ==================== 安全和加密工具 ====================

class SecurityTool(Tool):
    """安全和加密工具"""
    
    def __init__(self):
        super().__init__(
            name="security_tool",
            description="Handle encryption, authentication, and security operations"
        )
        self.fernet_key = None
        
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行安全操作"""
        try:
            if action == "encrypt":
                return self.encrypt_data(**kwargs)
            elif action == "decrypt":
                return self.decrypt_data(**kwargs)
            elif action == "hash":
                return self.hash_data(**kwargs)
            elif action == "generate_jwt":
                return self.generate_jwt(**kwargs)
            elif action == "verify_jwt":
                return self.verify_jwt(**kwargs)
            elif action == "generate_password":
                return self.generate_secure_password(**kwargs)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def encrypt_data(self, data: str, key: str = None) -> Dict[str, Any]:
        """加密数据"""
        try:
            if key:
                fernet = Fernet(key.encode() if isinstance(key, str) else key)
            else:
                # 生成新密钥
                key = Fernet.generate_key()
                fernet = Fernet(key)
                
            encrypted = fernet.encrypt(data.encode())
            
            return {
                "success": True,
                "encrypted_data": encrypted.decode(),
                "key": key.decode() if isinstance(key, bytes) else key
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def decrypt_data(self, encrypted_data: str, key: str) -> Dict[str, Any]:
        """解密数据"""
        try:
            fernet = Fernet(key.encode() if isinstance(key, str) else key)
            decrypted = fernet.decrypt(encrypted_data.encode())
            
            return {
                "success": True,
                "decrypted_data": decrypted.decode()
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def hash_data(self, data: str, algorithm: str = "sha256") -> Dict[str, Any]:
        """哈希数据"""
        try:
            if algorithm == "sha256":
                hash_obj = hashlib.sha256()
            elif algorithm == "sha512":
                hash_obj = hashlib.sha512()
            elif algorithm == "md5":
                hash_obj = hashlib.md5()
            else:
                return {"success": False, "error": f"Unsupported algorithm: {algorithm}"}
                
            hash_obj.update(data.encode())
            
            return {
                "success": True,
                "hash": hash_obj.hexdigest(),
                "algorithm": algorithm
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def generate_jwt(self, payload: Dict[str, Any], secret: str, 
                    algorithm: str = "HS256", expiry_hours: int = 24) -> Dict[str, Any]:
        """生成JWT令牌"""
        try:
            # 添加过期时间
            payload["exp"] = datetime.utcnow() + timedelta(hours=expiry_hours)
            payload["iat"] = datetime.utcnow()
            
            token = jwt.encode(payload, secret, algorithm=algorithm)
            
            return {
                "success": True,
                "token": token,
                "expires_in": expiry_hours * 3600
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def verify_jwt(self, token: str, secret: str, algorithm: str = "HS256") -> Dict[str, Any]:
        """验证JWT令牌"""
        try:
            payload = jwt.decode(token, secret, algorithms=[algorithm])
            
            return {
                "success": True,
                "valid": True,
                "payload": payload
            }
        except jwt.ExpiredSignatureError:
            return {"success": True, "valid": False, "error": "Token expired"}
        except jwt.InvalidTokenError as e:
            return {"success": True, "valid": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    def generate_secure_password(self, length: int = 16, 
                               include_symbols: bool = True) -> Dict[str, Any]:
        """生成安全密码"""
        import string
        import secrets
        
        try:
            characters = string.ascii_letters + string.digits
            if include_symbols:
                characters += string.punctuation
                
            password = ''.join(secrets.choice(characters) for _ in range(length))
            
            # 计算密码强度
            strength = "weak"
            if length >= 8 and any(c.islower() for c in password) and \
               any(c.isupper() for c in password) and any(c.isdigit() for c in password):
                strength = "medium"
                if include_symbols and any(c in string.punctuation for c in password):
                    strength = "strong"
                    
            return {
                "success": True,
                "password": password,
                "strength": strength
            }
        except Exception as e:
            return {"success": False, "error": str(e)}


# ==================== 通知和通信工具 ====================

class NotificationTool(Tool):
    """通知和通信工具"""
    
    def __init__(self):
        super().__init__(
            name="notification_tool",
            description="Send notifications via email, SMS, Slack, etc."
        )
        
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """执行通知操作"""
        try:
            if action == "send_email":
                return await self.send_email(**kwargs)
            elif action == "send_slack":
                return await self.send_slack_message(**kwargs)
            elif action == "send_webhook":
                return await self.send_webhook(**kwargs)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
            
    async def send_email(self, to: Union[str, List[str]], subject: str, 
                        body: str, attachments: List[str] = None,
                        smtp_config: Dict[str, Any] = None) -> Dict[str, Any]:
        """发送邮件"""
        try:
            # 使用默认配置或提供的配置
            config = smtp_config or {
                "host": os.getenv("SMTP_HOST", "smtp.gmail.com"),
                "port": int(os.getenv("SMTP_PORT", "587")),
                "username": os.getenv("SMTP_USERNAME"),
                "password": os.getenv("SMTP_PASSWORD"),
                "use_tls": True
            }
            
            msg = MIMEMultipart()
            msg['From'] = config['username']
            msg['To'] = ', '.join(to) if isinstance(to, list) else to
            msg['Subject'] = subject
            
            # 添加正文
            msg.attach(MIMEText(body, 'plain'))
            
            # 添加附件
            if attachments:
                for file_path in attachments:
                    with open(file_path, 'rb') as f:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(f.read())
                        encoders.encode_base64(part)
                        part.add_header(
                            'Content-Disposition',
                            f'attachment; filename= {Path(file_path).name}'
                        )
                        msg.attach(part)
                        
            # 发送邮件
            with smtplib.SMTP(config['host'], config['port']) as server:
                if config.get('use_tls'):
                    server.starttls()
                server.login(config['username'], config['password'])
                server.send_message(msg)
                
            return {
                "success": True,
                "message": "Email sent successfully",
                "recipients": to
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
